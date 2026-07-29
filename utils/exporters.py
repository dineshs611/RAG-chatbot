import io
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

def export_to_txt(chat_title, messages) -> bytes:
    """Export chat history to plain text format."""
    output = io.StringIO()
    output.write(f"--- Chat History: {chat_title} ---\n\n")
    
    for msg in messages:
        sender_label = "USER" if msg["sender"] == "user" else "AI ASSISTANT"
        output.write(f"[{msg['timestamp']}] {sender_label}:\n")
        output.write(f"{msg['text']}\n")
        
        if msg.get("citations"):
            output.write("Sources:\n")
            for cit in msg["citations"]:
                page_str = f", Page {cit.get('page')}" if cit.get('page') else ""
                output.write(f"  - {cit.get('document')}{page_str} (Conf: {int(cit.get('score', 0) * 100)}%)\n")
        output.write("-" * 50 + "\n\n")
        
    return output.getvalue().encode("utf-8")

def export_to_docx(chat_title, messages) -> bytes:
    """Export chat history to DOCX format using python-docx."""
    doc = Document()
    doc.add_heading(f"Chat History: {chat_title}", level=1)
    
    for msg in messages:
        sender_label = "Student" if msg["sender"] == "user" else "EduRAG AI Assistant"
        p = doc.add_paragraph()
        run = p.add_run(f"[{msg['timestamp']}] {sender_label}: ")
        run.bold = True
        
        if msg["sender"] == "user":
            run.font.color.rgb = None  # Default color or dark blue
        else:
            # AI colored dark purple
            pass
            
        doc.add_paragraph(msg["text"])
        
        if msg.get("citations"):
            p_cit = doc.add_paragraph()
            run_cit_title = p_cit.add_run("Sources Cited:\n")
            run_cit_title.italic = True
            run_cit_title.bold = True
            
            for cit in msg["citations"]:
                page_str = f", Page {cit.get('page')}" if cit.get('page') else ""
                doc.add_paragraph(
                    f"• {cit.get('document')}{page_str} (Confidence: {int(cit.get('score', 0) * 100)}%)",
                    style='List Bullet'
                )
        doc.add_paragraph("-" * 25)
        
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def export_to_pdf(chat_title, messages) -> bytes:
    """Export chat history to PDF using ReportLab."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54)
    story = []
    
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=20,
        leading=24,
        textColor=colors.HexColor('#7c3aed'),
        spaceAfter=20
    )
    
    sender_style_user = ParagraphStyle(
        'UserSender',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=10,
        leading=12,
        textColor=colors.HexColor('#2563eb'),
        spaceAfter=4
    )
    
    sender_style_ai = ParagraphStyle(
        'AiSender',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=10,
        leading=12,
        textColor=colors.HexColor('#7c3aed'),
        spaceAfter=4
    )
    
    msg_style = ParagraphStyle(
        'MsgBody',
        parent=styles['BodyText'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        textColor=colors.HexColor('#1f2937'),
        spaceAfter=8
    )
    
    citation_style = ParagraphStyle(
        'Citations',
        parent=styles['Normal'],
        fontName='Helvetica-Oblique',
        fontSize=8,
        leading=10,
        textColor=colors.HexColor('#4b5563'),
        leftIndent=15,
        spaceAfter=10
    )
    
    story.append(Paragraph(f"Chat History: {chat_title}", title_style))
    story.append(Spacer(1, 10))
    
    for msg in messages:
        sender_label = f"Student ({msg['timestamp']})" if msg["sender"] == "user" else f"EduRAG AI Assistant ({msg['timestamp']})"
        style_sender = sender_style_user if msg["sender"] == "user" else sender_style_ai
        
        story.append(Paragraph(sender_label, style_sender))
        
        # Replace newlines with break tags in reportlab paragraphs
        text_formatted = msg["text"].replace("\n", "<br/>")
        story.append(Paragraph(text_formatted, msg_style))
        
        if msg.get("citations"):
            cit_texts = []
            for cit in msg["citations"]:
                page_str = f", Page {cit.get('page')}" if cit.get('page') else ""
                cit_texts.append(f"• {cit.get('document')}{page_str} (Confidence: {int(cit.get('score', 0) * 100)}%)")
            
            citations_block = "<br/>".join(cit_texts)
            story.append(Paragraph(f"Sources cited:<br/>{citations_block}", citation_style))
            
        # Divider line using table
        t = Table([['']], colWidths=[500], rowHeights=[1])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#e5e7eb')),
            ('BOTTOMPADDING', (0,0), (-1,-1), 0),
            ('TOPPADDING', (0,0), (-1,-1), 0),
        ]))
        story.append(Spacer(1, 5))
        story.append(t)
        story.append(Spacer(1, 10))
        
    doc.build(story)
    return buffer.getvalue()
