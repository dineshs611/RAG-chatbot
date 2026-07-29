import io
import os
from docx import Document
from parsers.pdf import clean_text

def parse_docx(file_source):
    """
    Parse a Word document (.docx) from a path or raw bytes.
    Returns:
        List of dicts: [{'text': str, 'page_number': int}]
    """
    pages_data = []
    try:
        if isinstance(file_source, bytes):
            doc = Document(io.BytesIO(file_source))
        else:
            if not os.path.exists(file_source):
                raise FileNotFoundError(f"DOCX file not found: {file_source}")
            doc = Document(file_source)
            
        paragraphs_text = []
        for p in doc.paragraphs:
            if p.text.strip():
                paragraphs_text.append(p.text)
                
        # Also parse tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs_text.append(row_text)
                    
        # Simulate pages every 12 paragraphs/tables to support citation metadata
        paragraphs_per_page = 12
        num_simulated_pages = (len(paragraphs_text) + paragraphs_per_page - 1) // paragraphs_per_page
        
        for i in range(num_simulated_pages):
            chunk = paragraphs_text[i * paragraphs_per_page: (i + 1) * paragraphs_per_page]
            page_text = "\n\n".join(chunk)
            cleaned = clean_text(page_text)
            if cleaned:
                pages_data.append({
                    "text": cleaned,
                    "page_number": i + 1
                })
                
        if not pages_data:
            # Fallback if doc is empty but has headers/footers
            full_text = []
            for section in doc.sections:
                header = section.header
                for p in header.paragraphs:
                    if p.text.strip():
                        full_text.append(p.text)
            cleaned = clean_text("\n".join(full_text))
            if cleaned:
                pages_data.append({"text": cleaned, "page_number": 1})
                
    except Exception as e:
        print(f"Error parsing DOCX: {e}")
        raise e
        
    return pages_data
